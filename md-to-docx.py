import sys
import pathlib
import re
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


STYLE_NORMAL = 'Normal'
STYLE_NORMAL_COMPACT = 'Compact'
STYLE_DEFAULT_PARAGRAPH_TEXT = 'Body Text Char'
STYLE_HEADING1 = 'Heading 1'
STYLE_HEADING2 = 'Heading 2'
STYLE_HEADING3 = 'Heading 3'
STYLE_HEADING4 = 'Heading 4'
STYLE_HEADING5 = 'Heading 5'
STYLE_APPENDIX_HEADING1 = 'Kop 1 Bijlage'
STYLE_APPENDIX_HEADING2 = 'Kop 2 Bijlage'
STYLE_APPENDIX_HEADING3 = 'Kop 3 Bijlage'
STYLE_HEADING_TOC = 'Kop inhoudsopgave 1'
STYLE_LIST_BULLET = 'Lijst opsom.teken1'
STYLE_LIST_NUMBER = 'Lijstnummering1'
STYLE_TABLE = 'Tabelraster1'
STYLE_TITLE = 'Title'
STYLE_MAATREGEL = 'Maatregel'
STYLE_INSTRUCTION = 'Instructie'
DEFAULT_FONT = 'Calibri'
IMAGE_ICTU_LOGO = './Content/Images/ICTU.png'
IMAGE_WORDCLOUD = './Content/Images/word-cloud.png'
MD_BOLD = '**'
MD_BOLD_ALTERNATIVE = '__'
MD_ITALIC = '_'
MD_ITALIC_ALTERNATIVE = '*'
MD_STRIKETHROUGH = '~~'
MD_INSTRUCTION_START = '{'
MD_INSTRUCTION_END = '}'


def list_number(doc, par, prev=None, level=None, num=True):
    """
    This code was copied from: https://stackoverflow.com/questions/51829366/bullet-lists-in-python-docx

    Makes a paragraph into a list item with a specific level and
    optional restart.

    An attempt will be made to retreive an abstract numbering style that
    corresponds to the style of the paragraph. If that is not possible,
    the default numbering or bullet style will be used based on the
    ``num`` parameter.

    Parameters
    ----------
    doc : docx.document.Document
        The document to add the list into.
    par : docx.paragraph.Paragraph
        The paragraph to turn into a list item.
    prev : docx.paragraph.Paragraph or None
        The previous paragraph in the list. If specified, the numbering
        and styles will be taken as a continuation of this paragraph.
        If omitted, a new numbering scheme will be started.
    level : int or None
        The level of the paragraph within the outline. If ``prev`` is
        set, defaults to the same level as in ``prev``. Otherwise,
        defaults to zero.
    num : bool
        If ``prev`` is :py:obj:`None` and the style of the paragraph
        does not correspond to an existing numbering style, this will
        determine wether or not the list will be numbered or bulleted.
        The result is not guaranteed, but is fairly safe for most Word
        templates.
    """
    xpath_options = {
        True: {'single': 'count(w:lvl)=1 and ', 'level': 0},
        False: {'single': '', 'level': level},
    }

    def style_xpath(prefer_single=True):
        """
        The style comes from the outer-scope variable ``par.style.name``.
        """
        style = par.style.style_id
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:pStyle[@w:val="{style}"]'
            ']/@w:abstractNumId'
        ).format(style=style, **xpath_options[prefer_single])

    def type_xpath(prefer_single=True):
        """
        The type is from the outer-scope variable ``num``.
        """
        type = 'decimal' if num else 'bullet'
        return (
            'w:abstractNum['
                '{single}w:lvl[@w:ilvl="{level}"]/w:numFmt[@w:val="{type}"]'
            ']/@w:abstractNumId'
        ).format(type=type, **xpath_options[prefer_single])

    def get_abstract_id():
        """
        Select as follows:

            1. Match single-level by style (get min ID)
            2. Match exact style and level (get min ID)
            3. Match single-level decimal/bullet types (get min ID)
            4. Match decimal/bullet in requested level (get min ID)
            3. 0
        """
        for fn in (style_xpath, type_xpath):
            for prefer_single in (True, False):
                xpath = fn(prefer_single)
                ids = numbering.xpath(xpath)
                if ids:
                    return min(int(x) for x in ids)
        return 0

    if (prev is None or
            prev._p.pPr is None or
            prev._p.pPr.numPr is None or
            prev._p.pPr.numPr.numId is None):
        if level is None:
            level = 0
        numbering = doc.part.numbering_part.numbering_definitions._numbering
        # Compute the abstract ID first by style, then by num
        anum = get_abstract_id()
        # Set the concrete numbering based on the abstract numbering ID
        num = numbering.add_num(anum)
        # Make sure to override the abstract continuation property
        num.add_lvlOverride(ilvl=level).add_startOverride(1)
        # Extract the newly-allocated concrete numbering ID
        num = num.numId
    else:
        if level is None:
            level = prev._p.pPr.numPr.ilvl.val
        # Get the previous concrete numbering ID
        num = prev._p.pPr.numPr.numId.val
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
    par._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level


def format_paragraph(paragraph, line):
    processed_line, buffer = "", ""
    md_format_markers = (
        (MD_BOLD, MD_BOLD), (MD_BOLD_ALTERNATIVE, MD_BOLD_ALTERNATIVE),
        (MD_ITALIC, MD_ITALIC), (MD_ITALIC_ALTERNATIVE, MD_ITALIC_ALTERNATIVE),
        (MD_STRIKETHROUGH, MD_STRIKETHROUGH), (MD_INSTRUCTION_START, MD_INSTRUCTION_END))
    md_format_markers_to_keep = (MD_INSTRUCTION_START, MD_INSTRUCTION_END)

    def flush_buffer(markup=None):
        p = paragraph.add_run(text=buffer, style=STYLE_INSTRUCTION if markup == MD_INSTRUCTION_END else None)
        p.bold = markup in (MD_BOLD, MD_BOLD_ALTERNATIVE)
        p.italic = markup in (MD_ITALIC, MD_ITALIC_ALTERNATIVE)
        p.font.strike = markup == MD_STRIKETHROUGH

    while line:
        for md_format_start, md_format_end in md_format_markers:
            if line.startswith(md_format_start) and md_format_end in line[len(md_format_start):]:
                processed_chars = md_format_start
                flush_buffer()
                buffer = processed_chars if processed_chars in md_format_markers_to_keep else ''
                break
            elif line.startswith(md_format_end) and md_format_start in processed_line:
                processed_chars = md_format_end
                buffer += processed_chars if processed_chars in md_format_markers_to_keep else ''
                flush_buffer(md_format_end)
                buffer = ""
                break
        else:
            processed_chars = line[0]
            buffer += processed_chars
        line = line[len(processed_chars):]  # Move to the first character after the processed characters
        processed_line += processed_chars

    if buffer:
        paragraph.add_run(buffer)

    return paragraph


class ICTUDocument:
    """Wrap a docx Document and add ICTU specific styles and content."""

    def __init__(self, title, reference=None):
        self.document = Document(reference)
        if reference:
            self.clear_document()
        else:
            self.set_styles()
        self.create_header(title)
        self.create_frontpage(title)
        self.create_table_of_contents()

    def __getattr__(self, attr):
        return getattr(self.document, attr)  # Forward all unknown attribute lookups to the document

    def clear_document(self):
        for paragraph in self.document.paragraphs:
            self.delete_element(paragraph)
        for table in self.document.tables:
            self.delete_element(table)

    def set_styles(self):
        self.define_char_style(STYLE_DEFAULT_PARAGRAPH_TEXT, DEFAULT_FONT, Pt(12))
        self.define_style(STYLE_NORMAL, DEFAULT_FONT, Pt(12))
        self.define_style(STYLE_HEADING1, DEFAULT_FONT, Pt(32), top_margin=Pt(30), keep_with_next=True, page_break_before=True)
        self.define_style(STYLE_HEADING2, DEFAULT_FONT, Pt(18), top_margin=Pt(20), keep_with_next=True)
        self.define_style(STYLE_HEADING3, DEFAULT_FONT, Pt(12), top_margin=Pt(10), keep_with_next=True)
        self.define_style(STYLE_HEADING4, DEFAULT_FONT, Pt(12), keep_with_next=True)
        self.define_style(STYLE_HEADING5, DEFAULT_FONT, Pt(12), keep_with_next=True)

    def define_style(self, name, font_name, font_size, top_margin=Pt(0), keep_with_next=False, page_break_before=False):
        style = self.document.styles[name]
        style.font.name = font_name
        style.font.size = font_size
        style.paragraph_format.space_before = top_margin
        style.paragraph_format.keep_with_next = keep_with_next
        style.paragraph_format.page_break_before = page_break_before

    def define_char_style(self, name, font_name, font_size):
        style = self.document.styles[name]
        style.font.name = font_name
        style.font.size = font_size

    def create_header(self, title):
        self.document.sections[0].different_first_page_header_footer = True
        header_paragraph = self.document.sections[0].header.paragraphs[0]
        header_paragraph.clear()
        p = header_paragraph.add_run('{Projectnaam}')
        p.style = STYLE_INSTRUCTION
        header_paragraph.add_run(f'\t\t{title}')

    def create_frontpage(self, title):
        self.document.add_picture(IMAGE_ICTU_LOGO)
        self.document.add_paragraph(f'{title}', style=STYLE_TITLE)
        p = self.document.add_paragraph('', style=STYLE_NORMAL_COMPACT)
        p = p.add_run('{Projectnaam}')
        p.style = STYLE_INSTRUCTION
        p.bold = True
        self.document.add_paragraph('')
        p = self.document.add_paragraph('Versie ', style=STYLE_NORMAL_COMPACT)
        p = p.add_run('{versienummer}')
        p.style = STYLE_INSTRUCTION
        p = self.document.add_paragraph('Datum ', style=STYLE_NORMAL_COMPACT)
        p = p.add_run('{datum}')
        p.style = STYLE_INSTRUCTION
        for empty_paragraph in [''] * 4:
            self.document.add_paragraph(empty_paragraph)
        p = self.document.add_picture(IMAGE_WORDCLOUD)
        p.top_margin = Pt(36)
        self.document.add_page_break()

    def create_table_of_contents(self):
        self.document.add_paragraph("Inhoudsopgave", style=STYLE_HEADING_TOC)
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        for element in (fldChar, instrText, fldChar2, fldChar4):
            run._r.append(element)
        self.document.add_page_break()

    @staticmethod
    def delete_element(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None


class format_document:
    def __init__(self, input, output, reference, title):
        self.in_bijlagen = False
        self.in_list = False
        self.in_maatregel = False
        self.in_table = False
        styles = f"reference {reference}" if reference else "default styles"
        print(f"Converting {input} to {output} using {styles}.")
        self.document = ICTUDocument(title, reference)
        with open(input, mode='r', encoding='utf8') as source_file:
            lines = source_file.readlines()
        previous = None
        for line in lines:
            stripped_line = line.strip()
            if stripped_line:
                previous = self.process_line(stripped_line, previous, indented=line.startswith(' '))
        self.document.save(output)

    def process_line(self, line, previous, indented):
        p = None
        line, leaving_maatregel = self.process_maatregel_line(line)
        if line.startswith("|"):
            self.process_table_line(line)
        else:
            self.close_table()
            if line.startswith('#'):  # heading
                p = self.create_heading(line)
            elif re.match(r"^[*+-] .*", line):  # bullet list
                bullet_char = line[0]
                line = re.sub(r"^[*+-] ", "", line).strip()
                p = self.create_bullet_list_item(line, previous, bullet_char)
            elif re.match(r"^[0-9a-zA-Z]+[.] .*", line):  # numbered list
                level = 1 if line[0].isalpha() else (2 if indented else 0)
                line = re.sub(r"^[0-9a-zA-Z]+[.] ", "", line).strip()
                p = self.create_numbered_list_item(line, previous, level)
            else:  # regular text paragraph
                self.in_list = False
                p = self.create_paragraph(line)
        if leaving_maatregel:
            self.in_maatregel = False
        return p

    def process_maatregel_line(self, line):
        leaving_maatregel = False
        if line.startswith("@{"):
            self.in_maatregel = True
            line = line[2:]
        if self.in_maatregel and '}@' in line:
            leaving_maatregel = True
            line = line.replace('}@', '')
        return line, leaving_maatregel

    def process_table_line(self, line):
        if self.in_table:
            self.process_table_row(line)
        else:
            self.in_table = True
            self.process_header_row(line)

    def create_heading(self, line):
        heading_level = line.count('#')
        line = line.strip('#').strip()
        self.in_list = False
        if heading_level == 1 and line.upper() == 'BIJLAGEN':
            self.in_bijlagen = True
        if self.in_bijlagen:
            return self.document.add_paragraph(line, style=self.bijlage_heading_style(heading_level))
        else:
            return self.document.add_heading(line, level=heading_level)

    def create_numbered_list_item(self, line, previous, level=0):
        p = self.create_paragraph(line, style=STYLE_LIST_NUMBER)
        list_number(self.document, p, previous if self.in_list else None, level, num=True)
        self.in_list = True
        return p

    def create_bullet_list_item(self, line, previous, bullet_char):
        level = {'+': 1, '-': 2}.get(bullet_char, 0)
        p = self.create_paragraph(line, style=STYLE_LIST_BULLET)
        list_number(self.document, p, previous if self.in_list else None, level, num=False)
        self.in_list = True
        return p

    def create_paragraph(self, line, style=STYLE_NORMAL):
        p = self.document.add_paragraph("", style=STYLE_MAATREGEL if self.in_maatregel else style)
        return format_paragraph(p, line)

    def process_header_row(self, line):
        cells = self.get_cells(line)
        table = self.document.add_table(rows=1, cols=len(cells))
        table.style = STYLE_TABLE
        header_cells = table.rows[0].cells
        for header_cell, cell in zip(header_cells, cells):
            format_paragraph(header_cell.paragraphs[0], cell)

    def process_table_row(self, line):
        if len(self.document.tables[-1].rows) == 1 and '---' in line:
            self.process_table_row_alignment(line)
        else:
            self.process_table_row_content(line)

    def process_table_row_alignment(self, line):
        cell_alignment = []
        for cell in self.get_cells(line):
            if re.match(r"^:.*:$", cell):
                cell_alignment.append(WD_ALIGN_PARAGRAPH.CENTER)
            elif re.match(r".*:$", cell):
                cell_alignment.append(WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                cell_alignment.append(WD_ALIGN_PARAGRAPH.LEFT)
        for row in self.document.tables[-1].rows:
            for index, cell in enumerate(row.cells):
                cell.paragraphs[0].alignment = cell_alignment[index]

    def process_table_row_content(self, line):
        table = self.document.tables[-1]
        row_cells = table.add_row().cells
        for index, cell in enumerate(self.get_cells(line)):
            p = row_cells[index].paragraphs[0]
            format_paragraph(p, cell)
            p.alignment = table.rows[0].cells[index].paragraphs[0].alignment

    def close_table(self):
        if self.in_table:
            # Autofit is broken (https://github.com/python-openxml/python-docx/issues/209), use this work-around
            # instead of self.document.tables[-1].autofit = True
            for column in self.document.tables[-1].columns:
                for cell in column.cells:
                    cell._tc.get_or_add_tcPr().get_or_add_tcW().type = 'auto'
            self.in_table = False

    @staticmethod
    def get_cells(line):
        return [cell.strip() for cell in line.strip('| ').split('|')]

    @staticmethod
    def bijlage_heading_style(level):
        return {1: STYLE_APPENDIX_HEADING1, 2: STYLE_APPENDIX_HEADING2}.get(level, STYLE_APPENDIX_HEADING3)


if __name__ == "__main__":
    arguments = sys.argv[1:]
    if 3 <= len(arguments) <= 4:
        input, output, title, *optional_reference = arguments
        reference = optional_reference[0] if optional_reference else None
        format_document(input, output, reference, title)
    else:
        print(f"USAGE: md-to-docx <input file> <output file> <document title> [<reference file>]")
