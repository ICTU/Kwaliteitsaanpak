"""Docx builder unit tests."""

import pathlib
import tempfile
import unittest
from xml.etree.ElementTree import Element, ElementTree, SubElement

import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from pygments.util import ClassNotFound

import xmltags
from builder.docx_builder import DocxBuilder
from converter import Converter

REFERENCE_DOCX = pathlib.Path(__file__).parents[2] / "DocumentDefinitions" / "reference-neutral.docx"
BREAK = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"


class CodeBlockTest(unittest.TestCase):
    """Unit tests for code blocks in the docx builder."""

    def build(self, language: str, code: str) -> Document:
        """Build a docx with a single code block and return the resulting document."""
        document = Element(xmltags.DOCUMENT)
        code_block = SubElement(document, xmltags.CODE_BLOCK, {xmltags.CODE_BLOCK_LANGUAGE: language})
        code_block.text = code
        with tempfile.TemporaryDirectory() as directory:
            filename = pathlib.Path(directory) / "output.docx"
            Converter(ElementTree(document)).convert(DocxBuilder(filename, REFERENCE_DOCX, pathlib.Path(".")))
            return docx.Document(str(filename))

    def code_paragraph(self, document: Document) -> Paragraph:
        """Return the code block paragraph that was added (the last paragraph with the Code style)."""
        code_paragraphs = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.style is not None and paragraph.style.name == "Code"
        ]
        return code_paragraphs[-1]

    def test_code(self):
        """Test that the code is added verbatim, with line breaks between the lines."""
        paragraph = self.code_paragraph(self.build("python", "def f(x):\n    return x + 1\n"))
        self.assertEqual("def f(x):\n    return x + 1", paragraph.text)
        self.assertEqual(1, sum(len(run._r.findall(BREAK)) for run in paragraph.runs))

    def test_highlighting(self):
        """Test that tokens are highlighted, e.g. the keywords are green in the default style."""
        paragraph = self.code_paragraph(self.build("python", "def f(x):\n    return 1\n"))
        colors = {str(run.font.color.rgb) for run in paragraph.runs if run.font.color and run.font.color.rgb}
        self.assertIn("008000", colors)

    def test_yaml(self):
        """Test a YAML code block; the YAML lexer emits token types not styled in their own right."""
        paragraph = self.code_paragraph(self.build("yaml", "name: value\nlist:\n  - a\n"))
        self.assertEqual("name: value\nlist:\n  - a", paragraph.text)

    def test_unknown_language(self):
        """Test that an unknown language raises an exception."""
        self.assertRaises(ClassNotFound, self.build, "nonexistent", "code\n")


class TextFormattingTest(unittest.TestCase):
    """Unit tests for fomatting text."""

    def build(self, xmltag: str) -> Document:
        """Build an docx with a formatted paragraph."""
        document = Element(xmltags.DOCUMENT, {"title": "Title", "version": "1"})
        SubElement(document, xmltags.FRONTPAGE)  # Starts the <main> element that the document end closes
        paragraph = SubElement(document, xmltags.PARAGRAPH, {})
        formatted = SubElement(paragraph, xmltag)
        formatted.text = "formatted"
        with tempfile.TemporaryDirectory() as directory:
            filename = pathlib.Path(directory) / "output.docx"
            Converter(ElementTree(document)).convert(DocxBuilder(filename, REFERENCE_DOCX, pathlib.Path(".")))
            return docx.Document(str(filename))

    def test_bold(self):
        """Test bold text."""
        document = self.build(xmltags.BOLD)
        self.assertTrue(document.paragraphs[-1].runs[-1].bold)

    def test_italic(self):
        """Test italic text."""
        document = self.build(xmltags.ITALIC)
        self.assertTrue(document.paragraphs[-1].runs[-1].italic)

    def test_strikethrough(self):
        """Test strikethrough text."""
        document = self.build(xmltags.STRIKETHROUGH)
        self.assertTrue(document.paragraphs[-1].runs[-1].font.strike)

    def test_code(self):
        """Test code text."""
        document = self.build(xmltags.CODE)
        self.assertEqual("Fira code", document.paragraphs[-1].runs[-1].font.name)
