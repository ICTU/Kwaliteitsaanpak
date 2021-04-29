"""Docx builder."""

import pathlib
import shutil
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT
from docx.table import Table
from docx.table import _Row as Row
from docx.text.paragraph import Paragraph

import xmltags
from custom_types import TreeBuilderAttributes
from .builder import Builder
from .hyperlink import add_hyperlink
from .table_of_contents import add_table_of_contents


class DocxBuilder(Builder):
    """Docx builder."""

    SCHEMA = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    TEXT_TAGS = (
        xmltags.PARAGRAPH,
        xmltags.LIST_ITEM,
        xmltags.MEASURE,
        xmltags.HEADING,
        xmltags.TABLE_CELL,
        xmltags.HEADER,
        xmltags.TITLE,
        xmltags.INSTRUCTION,
        xmltags.BOLD,
        xmltags.ITALIC,
        xmltags.STRIKETHROUGH,
    )

    def __init__(self, filename: pathlib.Path, docx_reference_filename: pathlib.Path) -> None:
        super().__init__(filename)
        filename.unlink(missing_ok=True)
        shutil.copy(docx_reference_filename, filename)
        self.doc = Document(filename)
        self.paragraph: Optional[Paragraph] = None  # The current paragraph
        self.current_list_style: List[str] = []  # Stack of list styles
        self.previous_list_item: List[Optional[Paragraph]] = []  # Stack of previous list items
        self.table: Optional[Table] = None
        self.row: Optional[Row] = None
        self.column_index = 0

    def start_element(self, tag: str, attributes: TreeBuilderAttributes) -> None:
        # pylint: disable=protected-access
        super().start_element(tag, attributes)
        if tag == xmltags.PARAGRAPH:
            self.paragraph = self.doc.add_paragraph(style="Maatregel" if self.in_element(xmltags.MEASURE) else None)
        elif tag == xmltags.PAGEBREAK:
            self.doc.add_page_break()
        elif tag == xmltags.BULLET_LIST:
            self.current_list_style.append("Lijst opsom.teken1")
            self.previous_list_item.append(None)
        elif tag == xmltags.NUMBERED_LIST:
            self.current_list_style.append("Lijstnummering1")
            self.previous_list_item.append(None)
        elif tag == xmltags.LIST_ITEM:
            self._add_list_item()
        elif tag == xmltags.HEADING:
            level = self.nr_elements(xmltags.SECTION)
            in_appendix = self.in_element(xmltags.SECTION, {xmltags.SECTION_IS_APPENDIX: "y"})
            style = f"Kop {level} Bijlage" if in_appendix else f"heading {level}"
            self.paragraph = self.doc.add_paragraph(style=style)
        elif tag == xmltags.TABLE:
            self.table = self.doc.add_table(0, int(attributes[xmltags.TABLE_COLUMNS]), style="Tabelraster1")
            # Set table width to 100%
            self.table._tbl.tblPr.xpath("./w:tblW")[0].attrib[f"{self.SCHEMA}type"] = "pct"
            self.table._tbl.tblPr.xpath("./w:tblW")[0].attrib[f"{self.SCHEMA}w"] = "100%"
        elif tag in (xmltags.TABLE_HEADER_ROW, xmltags.TABLE_ROW):
            assert self.table
            self.row = self.table.add_row()
            self.column_index = 0
        elif tag == xmltags.TABLE_CELL:
            self._add_table_cell(attributes)
        elif tag == xmltags.HEADER:
            self.paragraph = self.doc.sections[0].header.paragraphs[0]
            self.paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # pylint: disable=no-member
        elif tag == xmltags.TITLE:
            self.paragraph = self.doc.add_paragraph(style="Title")
        elif tag == xmltags.TABLE_OF_CONTENTS:
            self.doc.add_paragraph(attributes[xmltags.TABLE_OF_CONTENTS_HEADING], style="TOC Heading")
            add_table_of_contents(self.doc.add_paragraph())
        elif tag == xmltags.IMAGE:
            self.doc.add_picture(attributes["src"][len("/work/") :])

    def _add_list_item(self) -> None:
        """Add a list item."""
        # pylint: disable=protected-access
        self.paragraph = self.doc.add_paragraph(style=self.current_list_style[-1])
        level = len(self.current_list_style) - 1
        self.paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_ilvl().val = level
        if self.current_list_style[-1] == "Lijstnummering1":
            if self.previous_list_item[-1] is None:
                # Add a new concrete numbering for Lijstnummering1. "0" is the id of the abstract numbering of
                # Lijstnummering1. This id can be found in the word/numbering.xml file (unzip reference.docx so
                # see word/numbering.xml), look for the <w:abstractNum w:abstractNumId="0" ...> that has a
                # child element <w:pStyle w:val="Lijstnummering1"/>
                num = self.doc.part.numbering_part.numbering_definitions._numbering.add_num("0")
                num.add_lvlOverride(ilvl=level).add_startOverride(1)  # Restart the numbering
                num = num.numId
            else:
                num = self.previous_list_item[-1]._p.pPr.numPr.numId.val
            self.paragraph._p.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = num
        self.previous_list_item[-1] = self.paragraph

    def _add_table_cell(self, attributes: TreeBuilderAttributes) -> None:
        """Add a table cell."""
        # pylint: disable=protected-access
        assert self.row
        cell = self.row.cells[self.column_index]
        cell._tc.tcPr.tcW.type = "auto"
        self.paragraph = cell.paragraphs[0]
        if alignment_attr := attributes.get(xmltags.TABLE_CELL_ALIGNMENT):
            # pylint: disable=no-member
            alignment = dict(
                left=WD_PARAGRAPH_ALIGNMENT.LEFT,
                right=WD_PARAGRAPH_ALIGNMENT.RIGHT,
                center=WD_PARAGRAPH_ALIGNMENT.CENTER,
            )[str(alignment_attr)]
            self.paragraph.paragraph_format.alignment = alignment
        self.column_index += 1

    def text(self, tag: str, text: str, attributes: TreeBuilderAttributes) -> None:
        if tag in self.TEXT_TAGS:
            assert self.paragraph
            run = self.paragraph.add_run(text)
            if self.in_element(xmltags.BOLD):
                run.font.bold = True
            if self.in_element(xmltags.INSTRUCTION):
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # pylint: disable=no-member
            if self.in_element(xmltags.ITALIC):
                run.font.italic = True
            if self.in_element(xmltags.STRIKETHROUGH):
                run.font.strike = True
        elif tag == xmltags.ANCHOR:
            assert self.paragraph
            try:
                link = str(attributes[xmltags.ANCHOR_LINK])
            except KeyError:
                print(tag, text, attributes)
                raise
            if link.startswith("#"):
                self.paragraph.add_run(text)  # Implement internal links some day
            else:
                add_hyperlink(self.paragraph, link, text)

    def end_element(self, tag: str, attributes: TreeBuilderAttributes) -> None:
        super().end_element(tag, attributes)
        if tag in (xmltags.BULLET_LIST, xmltags.NUMBERED_LIST):
            self.current_list_style.pop()
            self.previous_list_item.pop()

    def end_document(self) -> None:
        self.doc.save(self.filename)
